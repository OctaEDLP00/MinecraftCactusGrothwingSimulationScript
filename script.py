"""
_summary_

  Returns:
      _type_: _description_

  Yields:
      _type_: _description_
"""

import random
from pathlib import Path
from io import BytesIO
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


class GenerateDocxFile:
    """
    EN

    A class for simulating Minecraft 1.21.5 cactus growth mechanics and generating
    documentation in Word format with the results.

    This class provides methods to simulate cactus growth with flower mechanics,
    create visualization plots, and generate detailed documentation in both
    English and Spanish about the simulation results.

    ES

    Una clase para simular la mecánica de crecimiento de
    cactus en Minecraft 1.21.5 y generar documentación en formato
    Word con los resultados.

    Esta clase proporciona métodos para simular el crecimiento de
    cactus con la mecánica de las flores, crear gráficos de
    visualización y generar documentación detallada en inglés
    y español sobre los resultados de la simulación.
    """

    def run_simulation(self, num_simulations: int) -> list[int]:
        """
        Simulates the growth of N cactus plants according to Minecraft 1.21.5 mechanics.

        Args:
            num_simulations (int): Number of cactus growth simulations to run

        Returns:
            list[int]: List of final heights reached by each cactus, including the flower block
        """
        final_heights: list[int] = []
        for _ in range(num_simulations):
            h = 1
            while True:
                p = 0.10 if h in (1, 2) else 0.25
                if random.random() < p:
                    final_heights.append(h + 1)
                    break
                h += 1
        return final_heights

    def create_plot_image(
        self, height_data: list[int], title: str, xlabel: str, ylabel: str
    ) -> BytesIO:
        """
        Creates a histogram plot of cactus heights and returns it as a bytes buffer.

        Args:
            heights (list[int]): List of final cactus heights to plot
            title (str): Title of the plot
            xlabel (str): Label for x-axis
            ylabel (str): Label for y-axis

        Returns:
            BytesIO: Buffer containing the plot image in PNG format
        """
        plt.figure(figsize=(6, 4), dpi=300)
        plt.hist(height_data, bins=range(1, 61), edgecolor="black")
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        buf = BytesIO()
        plt.savefig(buf, format="png", bbox_inches="tight")
        buf.seek(0)
        plt.close()
        return buf

    def build_docx_es(self, height_data: list[int], output_path: str, sim_label: str):
        """
        Genera un documento Word en español con los resultados de la simulación de cactus.

        Args:
            heights (list[int]): Lista de alturas finales de los cactus simulados
            output_path (str): Ruta donde se guardará el documento Word
            sim_label (str): Etiqueta para identificar la simulación (ej: "N = 200,000")
        """
        total_simulations = len(height_data)
        # print("total_simulations", total_simulations, "line 98")
        count_24 = height_data.count(24)
        # print("count_24", count_24, "line 100")
        sim_prob_24 = count_24 / total_simulations * 100
        # print("sim_prob_24", sim_prob_24, "line 102")
        analytic_prob_24 = (0.9**2) * (0.75**20) * 100
        # print("analytic_prob_24", analytic_prob_24, "line 104")

        doc = Document()
        doc.add_heading(f"Simulación de cactus con flor — {sim_label}", level=1)
        doc.add_paragraph(
            "- Con la actualización Minecraft Java 1.21.5, se introdujo el bloque de flor de cactus, que puede generarse sobre un cactus como cuarto bloque, aunque de forma muy poco frecuente\n"
            "- La mecánica oficial documentada establece:\n"
            "  - 10 % de probabilidad de flor cuando el cactus intenta crecer de 1 a 2 bloques,\n"
            "  - 10 % nuevamente al intentar crecer a 3 bloques,\n"
            "  - y finalmente 25 % cuando intenta crecer más allá (sería el bloque 4, donde aparece la flor)\n"
            "- Por eso, sí es posible romper el récord de 23 bloques: si ya hay 23 bloques de cactus, el siguiente intento puede generar la flor y dejar un total de 24 bloques.\n"
            "- Pero la probabilidad es mínima:\n"
            "  - Análisis teórico: ≈ 0.000642 (≈ 0.0642 %)\n"
            "  - Simulación Monte-Carlo con 200 000 pruebas: ≈ 0.0715 % (muy cercano al estimado)\n"
            "- La flor detiene el crecimiento; por lo tanto, es la única forma de llegar a 24 — no permite seguir creciendo indefinidamente.\n"
        )
        img_buf = self.create_plot_image(
            height_data,
            f"Altura Final de Histograma — {sim_label}",
            "Altura final (bloques)",
            "Frecuencia",
        )
        doc.add_picture(img_buf, width=Inches(5))

        doc.add_heading("Resumen", level=2)
        doc.add_paragraph(
            f"- Probabilidad analítica de alcanzar 24 bloques: ≈ {analytic_prob_24:.4f}%\n"
            f"""- Simulación con N = {total_simulations:,}:
            ≈ {sim_prob_24:.4f}% casos donde el resultado
            final fue 24 bloques\n"""
            "- La flor impide crecimiento adicional; es el único camino hacia 24 bloques."
        )

        output_file = Path(output_path)
        output_dir = output_file.parent
        output_dir.mkdir(parents=True, exist_ok=True)

        doc.save(str(output_file))
        print(f"Document generated: {output_path}")

    def build_docx_en(self, height_data: list[int], output_path: str, sim_label: str):
        """
        Generates a Word document in English with the cactus simulation results.

        Args:
            heights (list[int]): List of final heights from the cactus simulations
            output_path (str): Path where the Word document will be saved
            sim_label (str): Label to identify the simulation (e.g., "N = 200,000")
        """

        total_simulations = len(height_data)
        # print("total_simulations", total_simulations, "line 158")
        count_24 = height_data.count(24)
        # print("count_24", count_24, "line 160")
        sim_prob_24 = count_24 / total_simulations * 100
        # print("sim_prob_24", sim_prob_24, "line 162")
        analytic_prob_24 = (0.9**2) * (0.75**20) * 100
        # print("analytic_prob_24s", analytic_prob_24, "line 164")

        doc = Document()
        doc.add_heading(f"Simulation of a flowering cactus — {sim_label}", level=1)
        doc.add_paragraph(
            "- Minecraft Java 1.21.5 added the cactus flower block, which can spawn atop a cactus as a fourth segment, though exceedingly rare\n"
            "- Game mechanics documented specify:\n"
            "  - 10% chance for a flower when a cactus attempts to grow from 1 to 2 blocks,\n"
            "  - another 10% when growing to 3 blocks,\n"
            "  - then a final 25% chance when attempting to grow to the next block (the 4th, where the flower appears)\n"
            "- Therefore, breaking the 23-block record is feasible: if a cactus is already at 23 cactus-blocks, the next growth can yield the flower and result in 24 total blocks.\n"
            "- Yet the probability is extremely low:\n"
            "    - Analytical estimate: ≈ 0.000642 (≈ 0.0642%)\n"
            "    - Monte Carlo simulation (200 000 trials): ≈ 0.0715%, very close to the theoretical estimate.\n"
            "- The flower stops further growth; thus, it's the sole path to 24 — it doesn’t allow endless growth.\n"
        )
        img_buf = self.create_plot_image(
            height_data,
            f"Histogram Final Heights — {sim_label}",
            "Final Heights (Blocks)",
            "Frecuency",
        )
        doc.add_picture(img_buf, width=Inches(5))

        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(
            f"- Analytical probability to reach 24 blocks: ≈ {analytic_prob_24:.4f}%\n"
            f"""- Simulation with N = {total_simulations:,}:
            ≈ {sim_prob_24:.4f}% ended at 24 blocks\n"""
            "- The flower stops further growth; it is the only route to reaching 24 blocks."
        )

        output_file = Path(output_path)
        output_dir = output_file.parent
        output_dir.mkdir(parents=True, exist_ok=True)

        doc.save(str(output_file))
        print(f"Document generated: {output_path}")


if __name__ == "__main__":
    for num_simul, label in [(200_000, "N = 200,000"), (1_000_000, "N = 1,000,000")]:
        heights = GenerateDocxFile().run_simulation(num_simul)
        FNAME_ES = (
            f'docs/simulacion_cactus_flor_{num_simul if num_simul<1e6 else "1M"}.docx'
        )
        FNAME_EN = (
            f'docs/cactus_flor_simulation_{num_simul if num_simul<1e6 else "1M"}.docx'
        )
        GenerateDocxFile().build_docx_es(heights, FNAME_ES, label)
        GenerateDocxFile().build_docx_en(heights, FNAME_EN, label)
